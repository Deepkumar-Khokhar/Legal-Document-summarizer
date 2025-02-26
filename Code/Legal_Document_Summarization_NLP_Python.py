#!/usr/bin/env python
# coding: utf-8

# In[2]:


# Student Information

# Student 1 Name: Belani hemil mayurkumar
# Student 2 Name: Nakrani kenil umeshbhai
# Student 3 Name: Khokhar deep kishorbhai
# Student 4 Name: Prajapati Aryan jitendrabhai
# Student 5 Name: Amin yesha pinakinbhai


# ## Installation of Required Libraries

# In[1]:


# !pip install transformers==2.8.0
# !pip install torch==1.4.0
# !pip install summa
# !pip install PyPDF2
# !pip install pdfplumber
# !pip install python-docx
# !pip install transformers torch summa rouge-score


# ## Importing Libraries
# 

# In[2]:


import os
import torch
import PyPDF2
import pdfplumber
from docx import Document
from transformers import T5Tokenizer, T5ForConditionalGeneration, T5Config
from transformers import BartForConditionalGeneration, BartTokenizer
from summa import summarizer


# ## Functions for Text Extraction

# In[3]:


def extract_text_from_pdf(file_path):
    text_array = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_array.append(text.replace('\n', ' '))
    return ' '.join(text_array)


# In[4]:


def extract_text_from_docx(file_path):
    text_array = []
    doc = Document(file_path)
    for para in doc.paragraphs:
        text = para.text.replace('\n', ' ')
        text_array.append(text)
    return ' '.join(text_array)


# In[5]:


file_name = 'free-software-development-agreement.docx'


# In[6]:


extension = os.path.splitext(file_name)[1].lower()


# In[7]:


if extension == '.pdf':
  text =  extract_text_from_pdf(file_name)
elif extension == '.docx':
  text =  extract_text_from_docx(file_name)
else:
    raise ValueError("Unsupported file type: " + extension)


# In[8]:


print(text)


# **T5 (Text-To-Text Transfer Transformer)**

# In[9]:


model = T5ForConditionalGeneration.from_pretrained('t5-small')
tokenizer = T5Tokenizer.from_pretrained('t5-small')
device = torch.device('cpu')


# In[10]:


preprocessed_text = text.strip().replace('\n','')
t5_input_text = 'summarize: ' + preprocessed_text


# In[11]:


tokenized_text = tokenizer.encode(t5_input_text, return_tensors='pt', max_length=512).to(device)


# In[12]:


summary_ids = model.generate(tokenized_text, min_length=30, max_length=200)
summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)


# In[13]:


print(summary)


# **BART Large (with Long Text Handling)**

# In[14]:


model_name = "facebook/bart-large-cnn"
tokenizer = BartTokenizer.from_pretrained(model_name)
model = BartForConditionalGeneration.from_pretrained(model_name)


# In[15]:


def summarize_text(text):
    inputs = tokenizer(text, return_tensors="pt", max_length=1024, truncation=True)
    summary_ids = model.generate(
        inputs['input_ids'],
        num_beams=4,
        max_length=200,
        min_length=40,
        length_penalty=2.0,
        early_stopping=True
    )
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary


# In[16]:


summary = summarize_text(text)
print(summary)


# **TextRank**

# In[17]:


def text_rank_summary(text, ratio=0.2):
    summary = summarizer.summarize(text, ratio=ratio)
    return summary


# In[18]:


summary = text_rank_summary(text, ratio=0.5)
print(summary)


# ## Model Comparison

# In[19]:


import torch
from transformers import T5Tokenizer, T5ForConditionalGeneration, BartForConditionalGeneration, BartTokenizer
from summa import summarizer
from rouge_score import rouge_scorer

# Define a sample reference summary (for evaluation purposes)
reference_summary = "This is a sample reference summary of a legal document for comparison purposes."

# Pre-load models and tokenizers
t5_model = T5ForConditionalGeneration.from_pretrained('t5-small')
t5_tokenizer = T5Tokenizer.from_pretrained('t5-small')

bart_model = BartForConditionalGeneration.from_pretrained('facebook/bart-large-cnn')
bart_tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')

device = torch.device('cpu')

# Summarization function for T5 model
def summarize_with_t5(text):
    t5_input_text = 'summarize: ' + text.strip().replace('\n', '')
    tokenized_text = t5_tokenizer.encode(t5_input_text, return_tensors='pt', max_length=512).to(device)
    summary_ids = t5_model.generate(tokenized_text, min_length=30, max_length=200)
    return t5_tokenizer.decode(summary_ids[0], skip_special_tokens=True)

# Summarization function for BART model
def summarize_with_bart(text):
    inputs = bart_tokenizer(text, return_tensors="pt", max_length=1024, truncation=True)
    summary_ids = bart_model.generate(
        inputs['input_ids'], num_beams=4, max_length=200, min_length=40, length_penalty=2.0, early_stopping=True
    )
    return bart_tokenizer.decode(summary_ids[0], skip_special_tokens=True)

# Summarization function for TextRank (Extractive)
def summarize_with_textrank(text):
    return summarizer.summarize(text, ratio=0.2)

# ROUGE scorer setup
scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)

# Evaluation function
def evaluate_summary(model_name, generated_summary):
    scores = scorer.score(reference_summary, generated_summary)
    print(f"\nModel: {model_name}")
    print(f"ROUGE-1: {scores['rouge1']}")
    print(f"ROUGE-2: {scores['rouge2']}")
    print(f"ROUGE-L: {scores['rougeL']}")

# Sample legal document text (for testing)
legal_document_text = "This is a sample legal document text to evaluate the summarization models. It contains several points related to legal agreements and clauses. The models will generate summaries based on this input."

# Get summaries from the three models
t5_summary = summarize_with_t5(legal_document_text)
bart_summary = summarize_with_bart(legal_document_text)
textrank_summary = summarize_with_textrank(legal_document_text)

# Evaluate and compare models using ROUGE
evaluate_summary("T5 Model", t5_summary)
evaluate_summary("BART Model", bart_summary)
evaluate_summary("TextRank Model", textrank_summary)


# ## GUI Part

# In[21]:


import ipywidgets as widgets
from IPython.display import display, HTML
from transformers import BartTokenizer, BartForConditionalGeneration
import os
import torch
import pdfplumber
from docx import Document

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    text_array = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_array.append(text.replace('\n', ' '))
    return ' '.join(text_array)

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    text_array = []
    doc = Document(file_path)
    for para in doc.paragraphs:
        text = para.text.replace('\n', ' ')
        text_array.append(text)
    return ' '.join(text_array)

# Summarization model (BART)
model = BartForConditionalGeneration.from_pretrained('facebook/bart-large-cnn')
tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')
device = torch.device('cpu')

# Function to summarize text
def summarize_text(text):
    preprocessed_text = text.strip().replace('\n', '')

    # Prepare input for BART
    inputs = tokenizer.encode("summarize: " + preprocessed_text, return_tensors='pt', max_length=1024, truncation=True).to(device)
    summary_ids = model.generate(inputs, min_length=30, max_length=200, length_penalty=2.0, num_beams=4, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

# Title label
title_label = widgets.HTML(
    value="<h2 style='text-align: center; color: #333;'>Legal Document Summarization Tool</h2>"
)

# File upload widget
file_upload = widgets.FileUpload(
    accept='.pdf,.docx',  # Accept only PDF and DOCX files
    multiple=False  # Only single file upload allowed
)

# Text area for displaying extracted content
text_area = widgets.Textarea(
    value='',
    placeholder='Extracted text will appear here...',
    description='Document Text:',
    layout=widgets.Layout(width='600px', height='300px', background_color='#f9f9f9'),
    style={'description_width': 'initial'}
)

# Text area for displaying the summary
summary_area = widgets.Textarea(
    value='',
    placeholder='Summary will appear here...',
    description='Summary:',
    layout=widgets.Layout(width='600px', height='200px', background_color='#f0f0f0'),
    style={'description_width': 'initial'}
)

# Function to handle file upload and text extraction
def on_file_upload(change):
    uploaded_file = list(change['new'].values())[0]
    file_name = uploaded_file['metadata']['name']
    file_content = uploaded_file['content']
    extension = os.path.splitext(file_name)[1].lower()

    try:
        # Save uploaded file temporarily
        with open(file_name, 'wb') as f:
            f.write(file_content)

        # Extract text based on file extension
        if extension == '.pdf':
            extracted_text = extract_text_from_pdf(file_name)
        elif extension == '.docx':
            extracted_text = extract_text_from_docx(file_name)
        else:
            extracted_text = "Unsupported file format."

        # Update the text area with the extracted content
        text_area.value = extracted_text

    except Exception as e:
        text_area.value = f"Error: {str(e)}"

# Function to summarize the extracted content when the summarize button is clicked
def on_summarize_button_clicked(b):
    document_text = text_area.value
    if document_text.strip():
        summary = summarize_text(document_text)
        summary_area.value = summary
    else:
        summary_area.value = "No text available for summarization."

# Button for summarizing the document (Centered using HBox)
summarize_button = widgets.Button(
    description="Summarize",
    button_style='success',  # Green button
    layout=widgets.Layout(width='150px')
)

# Center the button using an HBox
button_box = widgets.HBox([summarize_button], layout=widgets.Layout(justify_content='center'))

# Trigger the summarization process on button click
summarize_button.on_click(on_summarize_button_clicked)

# Trigger the text extraction on file upload
file_upload.observe(on_file_upload, names='value')

# Display widgets (title, file upload, extracted text area, summarize button, and summary area)
display(widgets.VBox([title_label, file_upload, text_area, button_box, summary_area]))


